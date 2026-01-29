"""
Production-Grade ATS-Optimized Resume Exporter

FEATURES:
- ATS-friendly DOCX formatting with standard fonts
- Clean, scannable layout without complex formatting
- Preserves section dividers as plain text
- Uses standard document structure for maximum compatibility
- Avoids tables, text boxes, and decorative elements that break ATS parsing

FONTS: Calibri (primary), Arial (fallback) - both ATS-friendly
STRUCTURE: Simple paragraphs and bullet points only
FORMATTING: Minimal styling to ensure ATS readability
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import black
from reportlab.lib.enums import TA_LEFT, TA_CENTER
import logging
import os

logger = logging.getLogger(__name__)


def export_to_docx(resume_text: str, output_path: str) -> None:
    """
    Export resume to ATS-optimized DOCX format
    
    Args:
        resume_text (str): Complete resume content
        output_path (str): File path for output document
    """
    try:
        logger.info(f"Exporting resume to DOCX: {output_path}")
        
        # Create new document
        doc = Document()
        
        # Set ATS-friendly margins (0.5-1 inch recommended)
        _set_document_margins(doc)
        
        # Process resume content line by line
        lines = resume_text.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            if not line:
                # Add minimal spacing for empty lines
                _add_paragraph(doc, "", space_after=3)
                continue
            
            # Determine line type and format accordingly
            line_type = _classify_line(line, i, lines)
            
            if line_type == 'name_header':
                _add_name_header(doc, line)
            elif line_type == 'contact_info':
                _add_contact_info(doc, line)
            elif line_type == 'section_header':
                _add_section_header(doc, line)
            elif line_type == 'separator':
                _add_separator(doc, line)
            elif line_type == 'bullet_point':
                _add_bullet_point(doc, line)
            elif line_type == 'sub_bullet':
                _add_sub_bullet_point(doc, line)
            else:
                _add_body_text(doc, line)
        
        # Save document
        doc.save(output_path)
        logger.info("Successfully exported DOCX resume")
        
    except Exception as e:
        logger.error(f"Failed to export DOCX: {e}")
        # Fallback: save as text file
        _save_as_text_fallback(resume_text, output_path)


def export_to_pdf(resume_text: str, output_path: str) -> None:
    """
    Export resume to ATS-optimized PDF format
    
    Args:
        resume_text (str): Complete resume content
        output_path (str): File path for output PDF
    """
    try:
        logger.info(f"Exporting resume to PDF: {output_path}")
        
        # Create PDF document with ATS-friendly settings
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch
        )
        
        # Create ATS-optimized styles
        styles = _create_pdf_styles()
        
        # Build PDF content
        story = []
        lines = resume_text.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            if not line:
                story.append(Spacer(1, 6))
                continue
            
            # Classify and format line
            line_type = _classify_line(line, i, lines)
            
            if line_type == 'name_header':
                story.append(Paragraph(line, styles['NameHeader']))
            elif line_type == 'contact_info':
                story.append(Paragraph(line, styles['ContactInfo']))
            elif line_type == 'section_header':
                story.append(Spacer(1, 8))
                story.append(Paragraph(line, styles['SectionHeader']))
            elif line_type == 'separator':
                # Skip separators in PDF (handled by spacing)
                continue
            elif line_type == 'bullet_point':
                bullet_text = _clean_bullet_text(line)
                story.append(Paragraph(f"• {bullet_text}", styles['BulletPoint']))
            elif line_type == 'sub_bullet':
                bullet_text = _clean_bullet_text(line)
                story.append(Paragraph(f"    • {bullet_text}", styles['SubBulletPoint']))
            else:
                story.append(Paragraph(line, styles['BodyText']))
        
        # Build PDF
        doc.build(story)
        logger.info("Successfully exported PDF resume")
        
    except Exception as e:
        logger.error(f"Failed to export PDF: {e}")
        # Fallback: save as text file
        _save_as_text_fallback(resume_text, output_path)


def _set_document_margins(doc: Document) -> None:
    """Set ATS-friendly document margins"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def _classify_line(line: str, index: int, all_lines: list) -> str:
    """Classify line type for appropriate formatting"""
    line_lower = line.lower().strip()
    
    # Name header (first significant line, all caps, not a section)
    if index <= 2 and line.isupper() and not _is_section_header(line):
        return 'name_header'
    
    # Contact information
    if any(indicator in line_lower for indicator in ['email:', 'phone:', '@', '|', '•']):
        if not line.startswith('•'):  # Not a bullet point
            return 'contact_info'
    
    # Section headers
    if _is_section_header(line):
        return 'section_header'
    
    # Separator lines
    if _is_separator_line(line):
        return 'separator'
    
    # Sub-bullet points (indented)
    if line.startswith('  •') or line.startswith('    •'):
        return 'sub_bullet'
    
    # Regular bullet points
    if line.startswith('•') or line.startswith('- '):
        return 'bullet_point'
    
    # Default to body text
    return 'body_text'


def _is_section_header(line: str) -> bool:
    """Check if line is a section header"""
    line_lower = line.lower().strip()
    
    # Known section headers
    section_headers = [
        'professional summary', 'professional objective', 'technical skills',
        'professional experience', 'education', 'projects', 'skills',
        'experience', 'summary', 'objective', 'certifications', 'awards',
        'languages', 'volunteer experience', 'publications', 'achievements'
    ]
    
    # Exact match or all caps short phrase
    return (line_lower in section_headers or 
            (line.isupper() and len(line.split()) <= 4 and len(line) > 3))


def _is_separator_line(line: str) -> bool:
    """Check if line is a separator"""
    return (len(line) > 2 and 
            all(c in '=-_' for c in line.strip()) and 
            len(set(line.strip())) <= 2)


def _add_name_header(doc: Document, text: str) -> None:
    """Add name header with ATS-friendly formatting"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(6)


def _add_contact_info(doc: Document, text: str) -> None:
    """Add contact information"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)


def _add_section_header(doc: Document, text: str) -> None:
    """Add section header with ATS-friendly formatting"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.bold = True
    p.space_before = Pt(12)
    p.space_after = Pt(3)


def _add_separator(doc: Document, text: str) -> None:
    """Add separator line (simplified for ATS)"""
    p = doc.add_paragraph()
    run = p.add_run('_' * 30)  # Simplified separator
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    p.space_after = Pt(6)


def _add_bullet_point(doc: Document, text: str) -> None:
    """Add bullet point with proper formatting"""
    bullet_text = _clean_bullet_text(text)
    p = doc.add_paragraph()
    run = p.add_run(f"• {bullet_text}")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    p.space_after = Pt(3)


def _add_sub_bullet_point(doc: Document, text: str) -> None:
    """Add sub-bullet point with indentation"""
    bullet_text = _clean_bullet_text(text)
    p = doc.add_paragraph()
    run = p.add_run(f"    • {bullet_text}")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    p.space_after = Pt(3)


def _add_body_text(doc: Document, text: str) -> None:
    """Add regular body text"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    p.space_after = Pt(6)


def _add_paragraph(doc: Document, text: str, space_after: int = 6) -> None:
    """Add paragraph with specified spacing"""
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    p.space_after = Pt(space_after)


def _clean_bullet_text(line: str) -> str:
    """Clean bullet point text by removing bullet characters"""
    cleaned = line.strip()
    
    # Remove various bullet characters
    if cleaned.startswith('•'):
        cleaned = cleaned[1:].strip()
    elif cleaned.startswith('- '):
        cleaned = cleaned[2:].strip()
    elif cleaned.startswith('  •'):
        cleaned = cleaned[3:].strip()
    elif cleaned.startswith('    •'):
        cleaned = cleaned[5:].strip()
    
    return cleaned


def _create_pdf_styles() -> dict:
    """Create ATS-optimized PDF styles"""
    styles = getSampleStyleSheet()
    
    custom_styles = {
        'NameHeader': ParagraphStyle(
            'NameHeader',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=8,
            spaceBefore=0,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ),
        'ContactInfo': ParagraphStyle(
            'ContactInfo',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica'
        ),
        'SectionHeader': ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=4,
            spaceBefore=8,
            alignment=TA_LEFT,
            fontName='Helvetica-Bold'
        ),
        'BodyText': ParagraphStyle(
            'BodyText',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            spaceBefore=0,
            alignment=TA_LEFT,
            fontName='Helvetica'
        ),
        'BulletPoint': ParagraphStyle(
            'BulletPoint',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=3,
            spaceBefore=0,
            leftIndent=0,
            alignment=TA_LEFT,
            fontName='Helvetica'
        ),
        'SubBulletPoint': ParagraphStyle(
            'SubBulletPoint',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=3,
            spaceBefore=0,
            leftIndent=20,
            alignment=TA_LEFT,
            fontName='Helvetica'
        )
    }
    
    return custom_styles


def _save_as_text_fallback(resume_text: str, output_path: str) -> None:
    """Save as text file if document export fails"""
    try:
        text_path = output_path.replace('.docx', '.txt').replace('.pdf', '.txt')
        with open(text_path, 'w', encoding='utf-8') as f:
            f.write(resume_text)
        logger.info(f"Saved fallback text file: {text_path}")
    except Exception as e:
        logger.error(f"Failed to save fallback text file: {e}")


# Legacy function names for backward compatibility
def export_resume_to_pdf(resume_text: str, output_path: str) -> None:
    """Legacy function - use export_to_pdf instead"""
    export_to_pdf(resume_text, output_path)


def export_resume_to_docx(resume_text: str, output_path: str) -> None:
    """Legacy function - use export_to_docx instead"""
    export_to_docx(resume_text, output_path)