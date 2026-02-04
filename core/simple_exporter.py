"""
Simple Resume Exporter - Handles DOCX and PDF export with fallbacks
"""

import os
import logging

logger = logging.getLogger(__name__)


def export_to_docx(resume_text, filepath):
    """Export resume text to DOCX file with multi-page support"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        
        logger.info(f"Creating DOCX file: {filepath}")
        logger.info(f"Resume text length: {len(resume_text)} characters")
        
        doc = Document()
        
        # Set margins for better multi-page layout
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Create custom styles
        styles = doc.styles
        
        # Name style - Black color, centered, 2pt font increase
        try:
            name_style = styles.add_style('ResumeName', WD_STYLE_TYPE.PARAGRAPH)
            name_font = name_style.font
            name_font.name = 'Calibri'
            name_font.size = Pt(13)  # 2pt increase from 11pt
            name_font.bold = True
            name_font.color.rgb = RGBColor(0, 0, 0)  # Black color
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centered
            name_style.paragraph_format.space_after = Pt(0)
            name_style.paragraph_format.space_before = Pt(0)
        except:
            name_style = None
        
        # Contact style - keep original, no changes
        try:
            contact_style = styles.add_style('ContactInfo', WD_STYLE_TYPE.PARAGRAPH)
            contact_font = contact_style.font
            contact_font.name = 'Calibri'
            contact_font.size = Pt(10)
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Keep original
            contact_style.paragraph_format.space_after = Pt(0)  # Keep original
            contact_style.paragraph_format.space_before = Pt(0)
        except:
            contact_style = None
        
        # Section header style - Only dark blue color, keep everything else original
        try:
            header_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
            header_font = header_style.font
            header_font.name = 'Calibri'  # Keep original font
            header_font.size = Pt(11)  # Keep original size
            header_font.bold = True
            header_font.color.rgb = RGBColor(30, 58, 138)  # Dark blue color only
            header_style.paragraph_format.space_before = Pt(0)  # Keep original spacing
            header_style.paragraph_format.space_after = Pt(0)
        except:
            header_style = None
        
        # Split resume into lines and process
        lines = resume_text.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            if not line:
                # Add spacing for empty lines
                doc.add_paragraph()
                continue
            
            # Check if it's a header (all caps or has equals/dashes)
            if line.isupper() and len(line.split()) <= 8 and not line.startswith('•') and not line.startswith('='):
                if '=' in line or '-' in line:
                    continue  # Skip separator lines
                
                # Check if this is the name (first line)
                if i == 0 or (i == 1 and lines[0].startswith('=')):
                    # This is the name
                    if name_style:
                        p = doc.add_paragraph(line, style='ResumeName')
                    else:
                        p = doc.add_paragraph(line)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.runs[0]
                        run.font.size = Pt(18)
                        run.font.bold = True
                else:
                    # Section header
                    if header_style:
                        p = doc.add_paragraph(line, style='SectionHeader')
                    else:
                        p = doc.add_paragraph()
                        run = p.add_run(line)
                        run.font.size = Pt(12)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 51, 102)
                        p.paragraph_format.space_before = Pt(12)
                        p.paragraph_format.space_after = Pt(6)
                        
            elif line.startswith('•') or line.startswith('-') or line.startswith('▸'):
                # Add as bullet point
                bullet_text = line[1:].strip() if line.startswith(('•', '-', '▸')) else line
                p = doc.add_paragraph(bullet_text, style='List Bullet')
                p.paragraph_format.space_after = Pt(3)
                
            elif line.startswith('  •') or line.startswith('  -'):
                # Sub-bullet point (indented)
                sub_bullet_text = line[3:].strip()
                p = doc.add_paragraph(sub_bullet_text, style='List Bullet 2')
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Inches(0.5)
                
            elif ':' in line and len(line) < 100 and ('Email:' in line or 'Phone:' in line or 'LinkedIn:' in line or 'Location:' in line):
                # Contact information
                if contact_style:
                    p = doc.add_paragraph(line, style='ContactInfo')
                else:
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.runs[0]
                    run.font.size = Pt(10)
                    p.paragraph_format.space_after = Pt(12)
                    
            else:
                # Add as regular paragraph
                # Handle very long paragraphs by splitting them
                if len(line) > 1000:
                    # Split long text into smaller chunks
                    words = line.split()
                    chunk_size = 150  # words per paragraph
                    for j in range(0, len(words), chunk_size):
                        chunk = ' '.join(words[j:j+chunk_size])
                        p = doc.add_paragraph(chunk)
                        p.paragraph_format.space_after = Pt(6)
                else:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(6)
        
        # Ensure directory exists
        dir_path = os.path.dirname(filepath)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)
        
        doc.save(filepath)
        logger.info(f"Successfully created DOCX: {filepath}")
        
        # Verify file was created and has content
        if os.path.exists(filepath) and os.path.getsize(filepath) > 0:
            logger.info(f"DOCX file size: {os.path.getsize(filepath)} bytes")
            return True
        else:
            raise Exception("DOCX file was not created or is empty")
        
    except ImportError as e:
        logger.error(f"Missing required library for DOCX export: {e}")
        return _save_as_text_fallback(resume_text, filepath, 'docx')
    except Exception as e:
        logger.error(f"Error creating DOCX: {e}")
        return _save_as_text_fallback(resume_text, filepath, 'docx')


def export_to_pdf(resume_text, filepath):
    """Export resume text to PDF file with multi-page support"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import darkblue, black, Color
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        
        # Define the exact teal/green color from web display
        teal_color = Color(45/255, 134/255, 89/255)  # #2d8659 converted to RGB
        
        logger.info(f"Creating PDF file: {filepath}")
        logger.info(f"Resume text length: {len(resume_text)} characters")
        
        # Ensure directory exists
        dir_path = os.path.dirname(filepath)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)
        
        # Create PDF document with proper margins for multi-page content
        doc = SimpleDocTemplate(
            filepath,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles for better formatting - matching web display
        name_style = ParagraphStyle(
            'ResumeName',
            parent=styles['Normal'],
            fontSize=16,  # Match web display (16pt)
            spaceAfter=4,
            spaceBefore=0,
            alignment=TA_CENTER,  # Centered
            textColor=black,  # Black color like web display
            fontName='Helvetica-Bold'
        )
        
        contact_style = ParagraphStyle(
            'ContactInfo',
            parent=styles['Normal'],
            fontSize=11,  # Match web display
            spaceAfter=4,
            spaceBefore=0,
            alignment=TA_CENTER,  # Centered like web display
            textColor=black
        )
        
        header_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Normal'],
            fontSize=12,  # Match web display (12pt)
            spaceAfter=4,
            spaceBefore=16,  # Match web display spacing
            textColor=teal_color,  # Exact teal/green color like web display
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'BodyText',
            parent=styles['Normal'],
            fontSize=11,  # Match web display (11pt)
            spaceAfter=0,  # Tight spacing like web display
            spaceBefore=0,
            leftIndent=0,
            textColor=black
        )
        
        bullet_style = ParagraphStyle(
            'BulletText',
            parent=styles['Normal'],
            fontSize=11,  # Match web display (11pt)
            spaceAfter=0,  # Tight spacing like web display
            spaceBefore=0,
            leftIndent=20,
            bulletIndent=10,
            textColor=black
        )
        
        subsection_style = ParagraphStyle(
            'SubsectionText',
            parent=styles['Normal'],
            fontSize=11,  # Match web display (11pt)
            spaceAfter=0,
            spaceBefore=3,  # Small spacing like web display
            textColor=black,
            fontName='Helvetica-Bold'  # Bold like web display
        )
        
        # Build content with proper page handling
        story = []
        lines = resume_text.split('\n')
        
        is_first_section = True
        current_section_lines = 0
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            if not line:
                # Add small spacer for empty lines
                story.append(Spacer(1, 2))
                continue
            
            # Skip dashed lines completely - they're just formatting artifacts
            if line.startswith('-') and len(set(line)) <= 2:
                continue
            
            # Detect section headers (should be teal/green with underline)
            section_headers = [
                'Professional Summary', 'Summary', 'Objective',
                'Skills', 'Technical Skills', 'Core Competencies', 
                'Education', 'Academic Background',
                'Experience', 'Professional Experience', 'Work Experience',
                'Projects', 'Key Projects', 'Notable Projects',
                'Certifications', 'Certificates', 'Awards',
                'Achievements', 'Accomplishments'
            ]
            
            if any(header.lower() in line.lower() for header in section_headers) and len(line) < 50:
                # Section header - add page break if needed for long sections
                if current_section_lines > 25 and not is_first_section:
                    story.append(PageBreak())
                    current_section_lines = 0
                
                # Create section header with simple underline (not a box)
                from reportlab.platypus import Table, TableStyle
                from reportlab.lib import colors
                
                # Create a simple paragraph with underline effect
                story.append(Paragraph(line, header_style))
                
                # Add a thin line underneath using a table
                line_table = Table([[''], ['']], colWidths=[6*inch], rowHeights=[0.01*inch, 0.05*inch])
                line_table.setStyle(TableStyle([
                    ('LINEBELOW', (0, 0), (0, 0), 1, teal_color),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(line_table)
                current_section_lines = 0
                is_first_section = False
                
            elif i == 0 or (line and not any(char in line for char in ['|', '@', 'http']) and len(line.split()) <= 4):
                # Name (first line or short line without contact info)
                if i == 0:
                    story.append(Paragraph(line, name_style))
                    current_section_lines = 0
                else:
                    # Could be a subsection header
                    story.append(Paragraph(line, subsection_style))
                    current_section_lines += 1
                    
            elif '|' in line and any(word in line.lower() for word in ['email', 'phone', 'linkedin', 'location', '@', 'http']):
                # Contact information
                story.append(Paragraph(line, contact_style))
                current_section_lines += 1
                
            elif line.startswith('•') or line.startswith('-') or line.startswith('▸'):
                # Bullet point
                bullet_text = line[1:].strip() if line.startswith(('•', '-', '▸')) else line
                story.append(Paragraph(f"• {bullet_text}", bullet_style))
                current_section_lines += 1
                
            elif line.startswith('  •') or line.startswith('  -'):
                # Sub-bullet point (indented)
                sub_bullet_text = line[3:].strip()
                sub_bullet_style = ParagraphStyle(
                    'SubBulletText',
                    parent=bullet_style,
                    leftIndent=40,
                    bulletIndent=30
                )
                story.append(Paragraph(f"• {sub_bullet_text}", sub_bullet_style))
                current_section_lines += 1
                
            else:
                # Check if it's a subsection header (job titles, education entries, etc.)
                if (len(line) < 100 and 
                    (any(word in line.lower() for word in ['university', 'college', 'bachelor', 'master', 'intern', 'engineer', 'developer', 'manager', 'graduated']) or
                     (',' in line and len(line.split(',')) == 2) or  # Job title, Company format
                     (len(line.split()) <= 10 and not line.startswith('•') and not line.startswith('-') and ':' not in line))):
                    # Likely a subsection header
                    story.append(Paragraph(line, subsection_style))
                else:
                    # Regular text
                    story.append(Paragraph(line, body_style))
                current_section_lines += 1
        
        # Build the PDF with automatic page breaks
        logger.info(f"Building PDF with {len(story)} elements")
        doc.build(story)
        logger.info(f"Successfully created PDF: {filepath}")
        
        # Verify file was created and has content
        if os.path.exists(filepath) and os.path.getsize(filepath) > 0:
            logger.info(f"PDF file size: {os.path.getsize(filepath)} bytes")
            return True
        else:
            raise Exception("PDF file was not created or is empty")
        
    except ImportError as e:
        logger.error(f"Missing required library for PDF export: {e}")
        return _save_as_text_fallback(resume_text, filepath, 'pdf')
    except Exception as e:
        logger.error(f"Error creating PDF: {e}")
        return _save_as_text_fallback(resume_text, filepath, 'pdf')


def _save_as_text_fallback(resume_text, original_filepath, original_format):
    """Save as text file when DOCX/PDF export fails"""
    try:
        # Create text file path
        text_filepath = original_filepath.replace(f'.{original_format}', '.txt')
        
        # Ensure directory exists
        dir_path = os.path.dirname(text_filepath)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)
        
        with open(text_filepath, 'w', encoding='utf-8') as f:
            f.write(resume_text)
        
        logger.info(f"Fallback: Saved as text file: {text_filepath}")
        
        # Update the original filepath to point to the text file
        if os.path.exists(text_filepath):
            # Copy the text file to the original path for download
            import shutil
            shutil.copy2(text_filepath, original_filepath)
            return True
        
        return False
        
    except Exception as e:
        logger.error(f"Failed to save fallback text file: {e}")
        return False


# Test function
def test_export():
    """Test export functionality"""
    test_resume = """JOHN DOE
========

Email: john@example.com | Phone: +1 555-123-4567

PROFESSIONAL SUMMARY
------------------------------
Experienced software engineer seeking new opportunities

TECHNICAL SKILLS
------------------------------
• Python
• JavaScript
• React

EXPERIENCE
------------------------------
• Software Engineer at Tech Company
• Developed web applications
"""
    
    print("Testing export functionality...")
    
    # Test DOCX export
    docx_result = export_to_docx(test_resume, "test_resume.docx")
    print(f"DOCX export: {'✅ Success' if docx_result else '❌ Failed'}")
    
    # Test PDF export
    pdf_result = export_to_pdf(test_resume, "test_resume.pdf")
    print(f"PDF export: {'✅ Success' if pdf_result else '❌ Failed'}")
    
    # Cleanup
    for file in ["test_resume.docx", "test_resume.pdf", "test_resume.txt"]:
        if os.path.exists(file):
            os.remove(file)


if __name__ == "__main__":
    test_export()    